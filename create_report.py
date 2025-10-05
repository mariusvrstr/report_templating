import pandas as pd
from pathlib import Path

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Cm, Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# ---------- PATHS ----------
script_dir = Path(__file__).resolve().parent
base_dir = script_dir / "data_source"
pictures_base_dir = base_dir / "pictures"

template_path = script_dir / "ReportTemplate.docx"
intermediate_path = script_dir / "CompletedReport_intermediate.docx"
output_path = script_dir / "CompletedReport.docx"

excel_file = base_dir / "ReportData.xlsx"

# ---------- HELPERS ----------
def load_kv_from_sheet(df: pd.DataFrame) -> dict:
    """If you use a 'General' sheet with 'Variable Name' / 'Variable Value'."""
    if df is None:
        return {}
    df = df.fillna("")
    return {row["Variable Name"]: row["Variable Value"] for _, row in df.iterrows()}

def format_headers(cols):
    out = []
    for c in cols:
        if isinstance(c, (pd.Timestamp, )):
            out.append(c.strftime("%d-%b-%y"))
        else:
            out.append(str(c))
    return out

def insert_table_at_marker(doc: Document, marker: str, df: pd.DataFrame,
                           widths_cm=None, table_style="Table Grid",
                           first_col_width_cm=8.0, other_col_width_cm=2.5):
    """
    Find a paragraph whose visible text EXACTLY matches `marker`,
    replace it with a real Word table built from df.
    """
    # default widths
    headers = format_headers(df.columns)
    if widths_cm is None:
        widths_cm = [first_col_width_cm] + [other_col_width_cm] * (len(headers) - 1)

    # build table at the end (python-docx limitation), then move it into place
    tbl = doc.add_table(rows=1, cols=len(headers))
    if table_style in [s.name for s in doc.styles]:
        tbl.style = table_style
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_cells = tbl.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h

    # data rows
    for _, row in df.fillna("").iterrows():
        cells = tbl.add_row().cells
        for j, val in enumerate(row.tolist()):
            cells[j].text = str(val)

    # set widths + alignment
    for col_idx, w in enumerate(widths_cm):
        for cell in tbl.columns[col_idx].cells:
            cell.width = Cm(w)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # center months, left-align first column
            for p in cell.paragraphs:
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # bold + center header
    for cell in tbl.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    # locate marker paragraph and replace it with the table
    def _replace_in_paragraphs(container):
        for para in container.paragraphs:
            if para.text.strip() == marker:
                # insert table xml directly after para
                para._element.addnext(tbl._element)
                # remove the marker paragraph
                para._element.getparent().remove(para._element)
                return True
        # also search inside tables (in case the marker was placed in a cell)
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    if _replace_in_paragraphs(cell):
                        return True
        return False

    found = _replace_in_paragraphs(doc)
    if not found:
        raise ValueError(f"Marker '{marker}' not found in document.")

# ---------- LOAD EXCEL ----------
excel = pd.read_excel(excel_file, sheet_name=None)

# optional: key/value variables from "General" sheet
general_ctx = load_kv_from_sheet(excel.get("General"))

# ---------- BUILD docxtpl CONTEXT ----------
context = {}
context.update(general_ctx)

# images (your existing multi-folder logic)
from docxtpl import DocxTemplate  # reuse below
doc_tpl = DocxTemplate(template_path)

picture_folders = [
    "approved_plan",
    "domestic_water_layout",
    "entrance",
    "sewer_layout",
    "storm_water_layout",
]
image_exts = (".png", ".jpg", ".jpeg")

for folder_name in picture_folders:
    folder_path = pictures_base_dir / folder_name
    var_name = f"{folder_name}_pictures"
    if not folder_path.exists():
        context[var_name] = None
        continue
    files = sorted([p for p in folder_path.glob("*") if p.suffix.lower() in image_exts],
                   key=lambda p: p.name.lower())
    imgs = []
    for p in files:
        try:
            imgs.append(InlineImage(doc_tpl, str(p), width=Mm(80)))
        except Exception:
            pass
    context[var_name] = imgs if imgs else None

# put literal markers into the rendered doc where tables should go
context["domestic_water_marker"] = "<<DOMESTIC_WATER_TABLE>>"
context["storm_water_marker"] = "<<STORM_WATER_TABLE>>"

# ---------- RENDER TEXT/IMAGES FIRST ----------
doc_tpl.render(context)
doc_tpl.save(intermediate_path)

# ---------- OPEN RESULT AND REPLACE MARKERS WITH REAL TABLES ----------
doc_final = Document(intermediate_path)

if "DomesticWater" in excel:
    df_dom = excel["DomesticWater"]
    insert_table_at_marker(
        doc_final,
        marker="<<DOMESTIC_WATER_TABLE>>",
        df=df_dom,
        # tweak widths if you want
        first_col_width_cm=8.0,
        other_col_width_cm=2.5,
        table_style="Table Grid",
    )

if "StormWater" in excel:
    df_storm = excel["StormWater"]
    insert_table_at_marker(
        doc_final,
        marker="<<STORM_WATER_TABLE>>",
        df=df_storm,
        first_col_width_cm=8.0,
        other_col_width_cm=2.5,
        table_style="Table Grid",
    )

doc_final.save(output_path)
print(f"✅ Report generated: {output_path}")
